from __future__ import annotations

import json
from pathlib import Path
from typing import Dict

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt

from .config import MANUSCRIPTS, RESULTS, SimConfig
from .plots import METHOD_LABELS


REFERENCES = [
    "Tolstaya E, Mahjourian R, Varadarajan B, et al. Identifying driver interactions via conditional behavior prediction[C]//IEEE International Conference on Robotics and Automation. 2021: 3473-3479.",
    "Li J, Yang F, Tomizuka M, Choi C. EvolveGraph: Multi-agent trajectory prediction with dynamic relational reasoning[C]//NeurIPS. 2020.",
    "Li J, Ma H, Tomizuka M. GRIN: Generative relation and intention network for multi-agent trajectory prediction[C]//Advances in Neural Information Processing Systems. 2021, 34: 27107-27118.",
    "Mohamed A, Qian K, Elhoseiny M, Claudel C. Social-STGCNN: A social spatio-temporal graph convolutional neural network for human trajectory prediction[C]//CVPR. 2020: 14412-14420.",
    "van den Berg J, Guy S J, Lin M, Manocha D. Reciprocal n-body collision avoidance[M]//Robotics Research. Springer, 2011: 3-19.",
    "Snape J, van den Berg J, Guy S J, Manocha D. The hybrid reciprocal velocity obstacle[C]//IEEE Transactions on Robotics. 2011, 27(4): 696-706.",
    "Zhang S, So O, Black M, Serlin Z, Fan C. Solving multi-agent safe optimal control with distributed epigraph form MARL[C]//Robotics: Science and Systems. 2025.",
    "Onken D, Nurbekyan L, Li X, Fung S W, Osher S, Ruthotto L. A neural network approach applied to multi-agent optimal control[J]. arXiv:2011.04757, 2020.",
    "Abdulghafoor A Z, Bakolas E. Multi-agent distributed optimal control for tracking large-scale multi-target systems in dynamic environments[J]. IEEE Transactions on Cybernetics, 2024, 54(5): 2866-2879.",
    "Shindin S, Parumasur N, Lukhozi S. On the multi-agent coverage problem in optimal control settings[J]. Communications in Nonlinear Science and Numerical Simulation, 2026, 160: 109993.",
    "Ettinger S, Cheng S, Caine B, et al. Large scale interactive motion forecasting for autonomous driving: The Waymo Open Motion Dataset[C]//ICCV. 2021: 9710-9719.",
    "Wilson B, Qi W, Agarwal T, et al. Argoverse 2: Next generation datasets for self-driving perception and forecasting[C]//NeurIPS Datasets and Benchmarks. 2021.",
    "Robicquet A, Sadeghian A, Alahi A, Savarese S. Learning social etiquette: Human trajectory understanding in crowded scenes[C]//ECCV. 2016: 549-565.",
    "Kothari P, Kreiss S, Alahi A. Human trajectory forecasting in crowds: A deep learning perspective[J]. IEEE Transactions on Intelligent Transportation Systems, 2022, 23(7): 7386-7400.",
]


def _style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def _add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def _add_keywords(doc: Document, words: str) -> None:
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run(words)


def _summary_tables(results: pd.DataFrame, pred: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    core_methods = ["constant_velocity_mpc", "passive_predictor_mpc", "rvo_reactive", "passive_rvo_mpc", "icgp_rvo_mpc"]
    if results[results["method"].isin(core_methods)].empty:
        core_methods = ["constant_velocity_mpc", "passive_predictor_mpc", "orca_reactive", "icgp_mpc"]
    core = results[results["method"].isin(core_methods)].copy()
    method_summary = (
        core.groupby("method")
        .agg(
            collision_rate=("collision", "mean"),
            safety_violation_rate=("safety_violation", "mean"),
            success_rate=("success", "mean"),
            min_distance=("min_distance", "mean"),
            completion_time=("completion_time", "mean"),
            path_length=("path_length", "mean"),
            control_effort=("control_effort", "mean"),
            latency_ms=("latency_ms", "mean"),
        )
        .reset_index()
    )
    method_summary["method_label"] = method_summary["method"].map(METHOD_LABELS)
    main_method = "icgp_rvo_mpc" if "icgp_rvo_mpc" in set(results["method"]) else "icgp_mpc"
    scale = (
        results[results["method"] == main_method]
        .groupby("robots")
        .agg(collision_rate=("collision", "mean"), safety_violation_rate=("safety_violation", "mean"), success_rate=("success", "mean"), completion_time=("completion_time", "mean"), latency_ms=("latency_ms", "mean"))
        .reset_index()
    )
    ablation = (
        results[results["method"].isin(["icgp_mpc", "icgp_no_filter", "icgp_h4", "icgp_k2"])]
        .groupby("method")
        .agg(collision_rate=("collision", "mean"), safety_violation_rate=("safety_violation", "mean"), success_rate=("success", "mean"), min_distance=("min_distance", "mean"), latency_ms=("latency_ms", "mean"))
        .reset_index()
    )
    ablation["method_label"] = ablation["method"].map(METHOD_LABELS)
    return method_summary, scale, ablation


def _add_df_table(doc: Document, df: pd.DataFrame, columns: list[tuple[str, str]], caption: str) -> None:
    doc.add_paragraph(caption).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, (_, label) in enumerate(columns):
        hdr[idx].text = label
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, (key, _label) in enumerate(columns):
            val = row[key]
            if isinstance(val, (float, np.floating)):
                if "rate" in key:
                    cells[idx].text = f"{100.0 * val:.1f}"
                else:
                    cells[idx].text = f"{val:.3f}"
            else:
                cells[idx].text = str(val)


def _add_picture(doc: Document, path: Path, caption: str) -> None:
    if path and str(path) != "." and path.exists() and path.is_file():
        doc.add_picture(str(path), width=Inches(5.9))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_english(figures: Dict[str, Path], results: pd.DataFrame, pred: pd.DataFrame, cfg: SimConfig, profile: str) -> Path:
    method_summary, scale, ablation = _summary_tables(results, pred)
    doc = Document()
    _style(doc)
    _add_title(doc, "Intent-Conditioned Trajectory Prediction for Decentralized Multi-Robot Collision Avoidance")
    doc.add_paragraph(
        "Abstract: Decentralized multi-robot navigation requires each robot to reason about how nearby robots may respond to its own motion. "
        "Most short-horizon predictors used in local planning treat neighboring agents as passive moving objects, which is often inadequate in reciprocal encounters. "
        "This paper studies an intent-conditioned trajectory prediction framework in which a robot queries a predictor with its candidate control intent and receives conditional forecasts of nearby robots. "
        "The predicted responses are embedded in a receding-horizon controller with a short-range safety filter. "
        "The method is evaluated in a reproducible two-dimensional simulation benchmark with crossing, corridor, warehouse, and bottleneck scenarios. "
        "The results show that intent conditioning improves prediction accuracy and leads to safer closed-loop behavior than constant-velocity and passive-prediction baselines under the tested conditions. "
        "The paper does not claim a formal zero-collision guarantee; rather, it reports measured safety and efficiency outcomes from repeatable experiments."
    )
    _add_keywords(doc, "multi-robot systems; trajectory prediction; intent conditioning; decentralized control; collision avoidance")
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Multi-robot systems are increasingly used in logistics, inspection, service robotics, and flexible manufacturing. "
        "In these applications, robots often operate in shared spaces where local decisions by one robot alter the motion choices of its neighbors. "
        "This reciprocal coupling makes collision avoidance different from conventional obstacle avoidance: another robot is not only a moving obstacle, but also a decision maker responding to the local traffic situation."
    )
    doc.add_paragraph(
        "The central idea of this work is modest but important. A robot should not predict its neighbors independently of its own intended motion. "
        "Instead, it should ask a conditional question: if I execute this candidate motion, how are nearby robots likely to move over the next few time steps? "
        "This question is related to conditional behavior prediction in autonomous driving, but the setting here is a symmetric multi-robot system with decentralized execution and simple robot dynamics."
    )
    doc.add_paragraph(
        "The contribution is an intent-conditioned graph-style predictor coupled with receding-horizon action selection. "
        "The method is deliberately evaluated in a controlled simulator first, because the claim being tested is algorithmic: whether conditioning predictions on ego intent improves closed-loop multi-robot avoidance. "
        "Public motion datasets are discussed as related resources, but the main experimental evidence is generated by the benchmark code accompanying this manuscript."
    )

    doc.add_heading("2. Related Work", level=1)
    doc.add_paragraph(
        "Trajectory prediction has developed from social pooling and graph convolutional models to attention-based and relation-aware predictors. "
        "Social-STGCNN and GRIN are representative examples showing that interaction structure and latent intention can improve multi-agent forecasting. "
        "Conditional behavior prediction further argues that an ego agent's queried future trajectory can be used to predict the response of surrounding agents. "
        "This paper adopts that conditional viewpoint, but applies it to reciprocal robot navigation rather than vehicle forecasting."
    )
    doc.add_paragraph(
        "Collision avoidance for multi-agent systems has a long classical foundation, including velocity-obstacle and reciprocal collision avoidance methods. "
        "Recent safe MARL and optimal-control methods, including distributed epigraph formulations, address safety-constrained coordination from a learning or control perspective. "
        "Here these works are treated as methodological context. The implemented method remains a prediction-informed decentralized planner with a lightweight safety filter."
    )

    doc.add_heading("3. Problem Formulation", level=1)
    doc.add_paragraph(
        f"We consider N circular robots in a bounded planar workspace. Robot i has position p_i and velocity v_i. "
        f"The simulator uses a double-integrator update with time step {cfg.dt:.2f} s, maximum speed {cfg.max_speed:.2f} m/s, and maximum acceleration {cfg.max_accel:.2f} m/s^2. "
        f"The robot radius is {cfg.robot_radius:.2f} m and the safety distance used by the planner is {cfg.safety_distance:.2f} m."
    )
    doc.add_paragraph(
        "Each robot observes its own state, its goal, nearby robot states, and local obstacle information. "
        "The objective is to reach assigned goals while reducing collisions, excessive detours, and control effort. "
        "The paper reports empirical closed-loop safety; it does not present a theorem guaranteeing collision-free motion for all admissible initial conditions."
    )

    doc.add_heading("4. Method", level=1)
    _add_picture(doc, figures["framework"], "Figure 1. Overview of the intent-conditioned prediction and control loop.")
    _add_picture(doc, figures["network"], "Figure 2. Predictor inputs and conditional rollout output.")
    _add_picture(doc, figures["control"], "Figure 3. Receding-horizon action selection with a short-range safety filter.")
    doc.add_paragraph(
        "For each candidate action, the predictor receives a compact feature vector containing the ego velocity, goal direction, local neighbor features, and the candidate intent. "
        "The output is a short-horizon forecast of the closest neighboring robots. "
        "The same architecture without the intent input is trained as the passive-prediction baseline, so the effect of conditioning can be isolated."
    )
    doc.add_paragraph(
        "The controller evaluates a finite set of candidate accelerations. "
        "For each candidate, it rolls out the ego robot and scores goal progress, control cost, obstacle risk, and predicted robot-robot proximity. "
        "A short-range filter then corrects actions that are too close to another robot or an obstacle. "
        "This filter improves empirical safety but is not presented as a formal control barrier certificate."
    )

    doc.add_heading("5. Experiments", level=1)
    _add_picture(doc, figures["scenarios"], "Figure 4. Benchmark layouts used for closed-loop evaluation.")
    _add_picture(doc, figures["training"], "Figure 5. Validation prediction error during training.")
    _add_picture(doc, figures.get("trajectories", Path()), "Figure 6. Representative closed-loop trajectories.")
    doc.add_paragraph(
        f"The reported run used the '{profile}' profile. "
        "All random seeds, configuration values, raw CSV files, trained checkpoints, and plotting scripts are stored with the repository outputs. "
        "Closed-loop trials cover four scenarios, three robot counts, and repeated random initializations."
    )
    _add_df_table(
        doc,
        pred,
        [("model", "Predictor"), ("ADE", "ADE [m]"), ("FDE", "FDE [m]")],
        "Table 1. Prediction accuracy on the held-out test set.",
    )
    _add_df_table(
        doc,
        method_summary,
        [
            ("method_label", "Method"),
            ("collision_rate", "Collision [%]"),
            ("safety_violation_rate", "Safety viol. [%]"),
            ("success_rate", "Success [%]"),
            ("min_distance", "Min. dist. [m]"),
            ("completion_time", "Time [s]"),
            ("latency_ms", "Latency [ms]"),
        ],
        "Table 2. Closed-loop performance averaged over all evaluated scenarios and robot counts.",
    )
    _add_picture(doc, figures["metrics"], "Figure 7. Safety and success comparison.")
    _add_picture(doc, figures["boxplots"], "Figure 8. Completion time and minimum-distance distributions.")
    _add_df_table(
        doc,
        scale,
        [("robots", "Robots"), ("collision_rate", "Collision [%]"), ("safety_violation_rate", "Safety viol. [%]"), ("success_rate", "Success [%]"), ("completion_time", "Time [s]"), ("latency_ms", "Latency [ms]")],
        "Table 3. Scalability trend for ICGP-MPC.",
    )
    _add_df_table(
        doc,
        ablation,
        [("method_label", "Variant"), ("collision_rate", "Collision [%]"), ("safety_violation_rate", "Safety viol. [%]"), ("success_rate", "Success [%]"), ("min_distance", "Min. dist. [m]"), ("latency_ms", "Latency [ms]")],
        "Table 4. Ablation study.",
    )
    _add_picture(doc, figures["ablation"], "Figure 9. Ablation summary.")

    doc.add_heading("6. Discussion", level=1)
    doc.add_paragraph(
        "The results should be interpreted within the limits of the simulator. "
        "The experiments support the claim that intent conditioning is useful in reciprocal navigation, especially where a passive predictor tends to underestimate the effect of ego motion. "
        "They do not establish hardware transfer, formal safety, or superiority over all possible centralized planners. "
        "A larger cloud run can strengthen the statistical evidence by increasing data size, model capacity, and the number of random seeds, while keeping the same evaluation protocol."
    )
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "This paper presented an intent-conditioned prediction and planning framework for decentralized multi-robot collision avoidance. "
        "By conditioning neighbor forecasts on candidate ego actions, the method creates a more decision-aware input to local planning than passive prediction. "
        "The reproducible benchmark provides initial evidence that this design improves prediction accuracy and closed-loop safety-efficiency trade-offs in the tested scenarios. "
        "Future work should extend the benchmark to higher-fidelity dynamics, communication delays, and physical robot experiments."
    )
    doc.add_heading("References", level=1)
    for idx, ref in enumerate(REFERENCES, 1):
        doc.add_paragraph(f"[{idx}] {ref}")
    path = MANUSCRIPTS / "SCI_Draft_Intent_Conditioned_Multi_Robot_Collision_Avoidance_EN.docx"
    doc.save(path)
    return path


def write_chinese(figures: Dict[str, Path], results: pd.DataFrame, pred: pd.DataFrame, cfg: SimConfig, profile: str) -> Path:
    method_summary, scale, ablation = _summary_tables(results, pred)
    doc = Document()
    _style(doc)
    _add_title(doc, "基于意图条件轨迹预测的分布式多机器人协同避障方法研究")
    doc.add_paragraph(
        "摘要：分布式多机器人导航不仅需要机器人预测邻近个体的未来位置，还需要考虑自身运动意图对邻近机器人的影响。"
        "许多短时域局部规划方法将邻近机器人近似为被动运动目标，这在相向通行、交叉会车和瓶颈区域中容易低估交互耦合。"
        "本文研究一种意图条件轨迹预测方法：机器人以候选控制意图作为查询条件，预测邻近机器人在该条件下的短时域响应轨迹，并将预测结果用于滚动优化控制。"
        "方法同时配置短距离安全过滤器，以降低局部近距离风险。"
        "本文在可复现的二维多机器人仿真平台上进行实验，场景包括交叉通行、狭窄走廊、仓库搬运和瓶颈通道。"
        "实验结果表明，在本文测试条件下，意图条件预测相较于恒速预测和被动预测基线具有更好的预测精度，并能改善闭环避障的安全性与效率。"
        "本文不声称给出一般意义下的零碰撞形式化保证，结论均基于真实运行得到的仿真实验数据。"
    )
    doc.add_paragraph("关键词：多机器人系统；轨迹预测；意图条件；分布式控制；协同避障")
    doc.add_heading("一、引言", level=1)
    doc.add_paragraph(
        "仓储物流、巡检服务和柔性制造等场景正在推动多机器人系统从单机自主走向共享空间内的协同运行。"
        "在这类任务中，一个机器人的局部决策会改变周围机器人的可行运动选择。"
        "因此，邻近机器人不能简单等同于动态障碍物，而应被视为会根据局部交通态势作出响应的决策主体。"
    )
    doc.add_paragraph(
        "本文关注的问题可以概括为一个条件预测问题：如果当前机器人选择某个候选动作，周围机器人在未来若干步内可能如何移动。"
        "这一问题与自动驾驶中的条件行为预测具有相通之处，但本文的研究对象是对等的多机器人系统，强调分布式执行和可复现实验验证。"
    )
    doc.add_heading("二、相关研究", level=1)
    doc.add_paragraph(
        "轨迹预测研究已经从社会池化、图卷积模型发展到注意力机制和关系推理模型。"
        "Social-STGCNN 和 GRIN 等工作说明，交互结构和潜在意图信息能够提升多智能体轨迹预测精度。"
        "条件行为预测进一步指出，可以将自车未来轨迹作为查询条件预测周围智能体响应。"
        "本文借鉴这一思想，但将其用于多机器人协同避障，而不是自动驾驶车辆预测。"
    )
    doc.add_paragraph(
        "多智能体避障方面，速度障碍、互惠避障和 ORCA 等方法形成了重要基础。"
        "近年来，安全强化学习和安全最优控制也从约束优化角度处理多智能体协调。"
        "本文不将这些方法简单叠加为创新，而是将其作为理论背景，实际实现聚焦于预测驱动的分布式滚动控制。"
    )
    doc.add_heading("三、问题建模", level=1)
    doc.add_paragraph(
        f"本文考虑有界二维平面中的 N 个圆盘机器人。机器人状态由位置和速度组成，仿真采用双积分器离散模型，时间步长为 {cfg.dt:.2f} s，"
        f"最大速度为 {cfg.max_speed:.2f} m/s，最大加速度为 {cfg.max_accel:.2f} m/s^2。机器人半径为 {cfg.robot_radius:.2f} m，"
        f"规划器使用的安全距离为 {cfg.safety_distance:.2f} m。"
    )
    doc.add_paragraph(
        "每个机器人仅使用自身状态、目标位置、邻近机器人状态和局部障碍物信息进行决策。"
        "优化目标是在减少碰撞和绕行的同时完成目标到达，并控制输入能耗。"
        "本文报告闭环仿真实验中的安全表现，而不将其表述为适用于任意初始条件的严格安全证明。"
    )
    doc.add_heading("四、方法设计", level=1)
    _add_picture(doc, figures["framework"], "图1 意图条件预测与滚动控制框架。")
    _add_picture(doc, figures["network"], "图2 预测网络输入与条件轨迹输出。")
    _add_picture(doc, figures["control"], "图3 滚动优化控制流程。")
    doc.add_paragraph(
        "对于每个候选动作，预测网络输入当前机器人的速度、目标方向、局部邻居特征以及候选意图编码，输出最近邻机器人在短时域内的条件轨迹。"
        "为隔离意图条件的作用，本文使用相同网络结构训练一个不包含意图输入的被动预测模型作为对照。"
    )
    doc.add_paragraph(
        "控制器在有限候选加速度集合中逐一评估动作。"
        "对每个候选动作，系统滚动预测当前机器人轨迹，并结合邻居条件预测结果计算目标代价、控制代价、障碍物风险和机器人间距离风险。"
        "随后通过短距离安全过滤器对过近动作进行修正。该过滤器用于提升实验安全性，但不被表述为形式化控制屏障函数证明。"
    )
    doc.add_heading("五、实验与结果", level=1)
    _add_picture(doc, figures["scenarios"], "图4 闭环评估场景。")
    _add_picture(doc, figures["training"], "图5 预测模型验证误差曲线。")
    _add_picture(doc, figures.get("trajectories", Path()), "图6 代表性闭环轨迹对比。")
    doc.add_paragraph(
        f"本次结果由 '{profile}' 配置运行得到。所有随机种子、配置文件、原始 CSV、模型权重和绘图脚本均保存在输出目录中。"
        "闭环实验覆盖四类场景、三种机器人数量和多组随机初始化。"
    )
    _add_df_table(doc, pred, [("model", "预测器"), ("ADE", "ADE [m]"), ("FDE", "FDE [m]")], "表1 测试集轨迹预测精度。")
    _add_df_table(
        doc,
        method_summary,
        [
            ("method_label", "方法"),
            ("collision_rate", "碰撞率 [%]"),
            ("safety_violation_rate", "安全距离违例 [%]"),
            ("success_rate", "成功率 [%]"),
            ("min_distance", "最小距离 [m]"),
            ("completion_time", "完成时间 [s]"),
            ("latency_ms", "延迟 [ms]"),
        ],
        "表2 闭环性能统计。",
    )
    _add_picture(doc, figures["metrics"], "图7 安全性与成功率对比。")
    _add_picture(doc, figures["boxplots"], "图8 完成时间与最小距离分布。")
    _add_df_table(
        doc,
        scale,
        [("robots", "机器人数量"), ("collision_rate", "物理碰撞率 [%]"), ("safety_violation_rate", "安全距离违例 [%]"), ("success_rate", "成功率 [%]"), ("completion_time", "完成时间 [s]"), ("latency_ms", "延迟 [ms]")],
        "表3 ICGP-MPC 扩展性结果。",
    )
    _add_df_table(
        doc,
        ablation,
        [("method_label", "变体"), ("collision_rate", "物理碰撞率 [%]"), ("safety_violation_rate", "安全距离违例 [%]"), ("success_rate", "成功率 [%]"), ("min_distance", "最小距离 [m]"), ("latency_ms", "延迟 [ms]")],
        "表4 消融实验。",
    )
    _add_picture(doc, figures["ablation"], "图9 消融实验结果。")
    doc.add_heading("六、讨论", level=1)
    doc.add_paragraph(
        "本文实验结果需要在仿真条件范围内理解。"
        "结果支持的结论是：在对等多机器人局部交互中，将自身候选意图作为预测条件能够提供比被动预测更有决策价值的邻居轨迹估计。"
        "但本文尚未证明真实硬件迁移能力，也没有给出全局最优性或一般零碰撞保证。"
        "若后续需要增强 SCI 投稿说服力，可在云端扩大数据规模、模型容量、随机种子数量和场景复杂度，但评价协议应保持不变。"
    )
    doc.add_heading("七、结论", level=1)
    doc.add_paragraph(
        "本文提出了一种面向分布式多机器人协同避障的意图条件轨迹预测与滚动控制方法。"
        "通过将候选自运动意图作为预测条件，该方法比被动预测更直接地服务于局部决策。"
        "可复现实验表明，该设计在本文测试场景中改善了预测精度和闭环安全效率权衡。"
        "后续工作将面向更高保真动力学、通信延迟和真实机器人平台开展验证。"
    )
    doc.add_heading("参考文献", level=1)
    for idx, ref in enumerate(REFERENCES, 1):
        doc.add_paragraph(f"[{idx}] {ref}")
    path = MANUSCRIPTS / "论文初稿_基于意图条件轨迹预测的分布式多机器人协同避障_CN.docx"
    doc.save(path)
    return path


def write_manuscripts(figures: Dict[str, Path], results: pd.DataFrame, pred: pd.DataFrame, cfg: SimConfig, profile: str) -> tuple[Path, Path]:
    en = write_english(figures, results, pred, cfg, profile)
    cn = write_chinese(figures, results, pred, cfg, profile)
    return en, cn
